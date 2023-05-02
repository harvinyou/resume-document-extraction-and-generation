import os
import xlrd
import datetime
import time
from mailmerge import MailMerge
import docx
import pickle
from docx import Document #导入库
import numpy as np
import json
import os
import json
import re
from io import StringIO
import xlrd
from py2neo import Graph, Node, Relationship
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from openpyxl import Workbook
from openpyxl.workbook import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font
modeljson='输出pdf抽取的信息_使用模型.json'
rulejson='输出pdf抽取的信息_基于规则.json'
truejson='验证集.json'
xlsxpath ='验证集.xlsx'
jsonpath='验证集.json'
path = 'valdocx' #文件路径
word_path='valdocx'
outpath = '文档填报普通表格'
excelp=xlrd.open_workbook(xlsxpath)
excelptable = excelp.sheet_by_name(excelp.sheet_names()[0])
nrows = excelptable.nrows
listtt1=['姓名', '出生年月', '性别', '电话', '籍贯', '落户市县', '政治面貌', '学位', '毕业时间', '工作时间', '项目时间', '毕业院校', '工作单位',
                 '工作内容', '职务', '项目名称', '项目责任']
listdate=[]


list1=['姓名', '出生年月', '性别', '电话', '籍贯', '落户市县', '政治面貌']
listt=['毕业院校', '工作单位','工作内容', '职务', '项目名称', '项目责任']#不能改
listt1=[ '学位', '毕业时间', '工作时间', '项目时间']
listtt1=['姓名', '出生年月', '性别', '电话', '籍贯', '落户市县', '政治面貌', '学位', '毕业时间', '工作时间', '项目时间', '毕业院校', '工作单位',
                 '工作内容', '职务', '项目名称', '项目责任']
listt2=['毕业院校', '工作单位','工作内容', '职务', '项目名称', '项目责任']


for i in excelptable .row_values(0):
    listdate.append(i)
print(listdate)
listdate=listtt1
def counttrain(word_path,jsonpath,key):
    count=0
    with open(jsonpath, 'r', encoding='utf-8') as j:
        truejson_info = json.load(j)
        val_paths = os.listdir(word_path)
        for p in val_paths:
            if p.endswith('.docx'):
                val_filename = p[:-5]
                num = num + 1
                info1 = truejson_info[val_filename]
                if info1!='':
                    try:

                        if isinstance(info1[key], str):
                            typeflag = 0
                            if key in info1:
                                count += 1
                        if isinstance(info1[key], list):
                            typeflag = 1
                            if key in info1:
                                count += 1
                    except Exception as e:
                        f = 1
    return count




def hasdate(content):
    for i in listdate:
        if i in cleanstr(content):
            return i
    return ''
def ext_general_field(content):#取右边第一个元素作为value
    con_len = len(content)
    for i, c in enumerate(content):
        if  i < con_len-1:
           return content[i+1]
    return ''
def cleanstr(c):
    listc = list(c)
    while ' ' in listc:
        listc.remove(' ')
    while '\n' in listc:
        listc.remove('\n')
    while '\xad' in listc:
        listc.remove('\xad')
    return ''.join(listc)

def wordpick(word_path,outpath,jsonpath):
    with open(jsonpath, 'r', encoding='utf-8') as j:
        truejson_info = json.load(j)
    files_list = os.listdir(word_path)
    tflag=0
    for file_name in files_list:
        # 判断文件类型是否在文件名中
        if '.docx' in file_name:
            wordname=file_name[:-5]#获取名字
            path1 = word_path + "\\"+file_name
            document=docx.Document(path1)
            tables = document.tables  # 获取文件中的表格集
            if wordname in truejson_info:
                tflag=1
            path_name = os.path.join(outpath, datetime.datetime.now().strftime("%Y-%m-%d"))
            if not os.path.exists(path_name):
                os.makedirs(path_name)
            word_name2 = path_name + "\\" + wordname + '.docx'
            if tflag == 1:
                document.save(word_name2)
                print(word_name2)

def wordtotest(word_path,outpath,jsonpath,key):
    w=0
    truecount = 0
    testcount = 0
    with open(jsonpath, 'r', encoding='utf-8') as j:
        truejson_info = json.load(j)
    files_list = os.listdir(word_path)
    for file_name in files_list:
        namelist=[]
        namelist2=[]
        # 判断文件类型是否在文件名中
        if '.docx' in file_name:
            wordname=file_name[:-5]#获取名字
            path1 = word_path + "\\"+file_name
            document=docx.Document(path1)
            tables = document.tables  # 获取文件中的表格集
            tflag=0
            if wordname in truejson_info:
                print(wordname)
                info1 = truejson_info[wordname]
                if key in info1:
                    truecount += 1

                for table in tables[:]:
                    for i, row in enumerate(table.rows[:]):  # 读每行
                        row_content = []
                        for j, cell in enumerate(row.cells[:]):  # 读一行中的所有单元格
                            rowmax=len(row.cells[:])
                            c = cell.text  # c是每一个单元格的字符串
                            if c == '':
                                flag = 1
                            else:
                                flag = 0
                                if  key in info1:
                                    if key in hasdate(c) or hasdate(c) in key:#检索
                                        tflag=1
            if tflag==1:
                testcount += 1
            print(key,truecount,testcount,file_name)
    return(truecount,testcount,testcount/truecount)





def allscore(list1):
    listt=[]
    listv=[]
    listw=[]
    listp=[]
    pweighted = 0
    allw=0
    for i in list1:
        key = i
        truecount,testcount,p=wordtotest(word_path,outpath,jsonpath,key)
        listt.append(truecount)
        listv.append(testcount)
        w=truecount
        listp.append(p)
        listw.append(w)
        allw+=truecount
        pweighted+=p*w
    pw=pweighted/allw



    result_wb = Workbook()
    # 第一个sheet是ws
    ws1 = result_wb.worksheets[0]
    # ws1=wb1.create_sheet('result',0)
    # 设置ws的名称
    ws1.title = "简历"

    ft = Font(name='Arial', size=11, bold=True)
    for k in range(len(list1)):
        ws1.cell(row=1, column=k + 2).value = list1[k]
        ws1.cell(row=1, column=k + 2).font = ft

    j = 2
    for p in listp:
        ws1.cell(row=2, column=1).value = '召回率'
        ws1.cell(row=2, column=j).value = float(p)
        j = j + 1
    j = 2
    for p in listt:
        ws1.cell(row=3, column=1).value = '原文数'
        ws1.cell(row=3, column=j).value = float(p)
        j = j + 1
    j = 2
    for w in listw:
        ws1.cell(row=4, column=1).value = '权重'
        ws1.cell(row=4, column=j).value = float(w/ allw)
        j = j + 1


    ws1.cell(row=5, column=1).value = '加权召回率'
    ws1.cell(row=5, column=2).value = float(pw)

    result_wb.save(filename='结果601.xlsx')


#wordtotest(word_path,outpath,jsonpath,'姓名')

allscore(listdate)
#wordpick(word_path,outpath,modeljson)
